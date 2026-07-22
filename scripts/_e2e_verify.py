"""End-to-end integration verification against a running API."""
from __future__ import annotations

import io
import json
import time
import uuid
from pathlib import Path
from typing import Any

import httpx

BASE = "http://127.0.0.1:8000/api/v1"
REPORT: list[dict[str, Any]] = []


def record(area: str, name: str, ok: bool, detail: str = "", ms: float | None = None) -> None:
    REPORT.append(
        {
            "area": area,
            "name": name,
            "ok": ok,
            "detail": detail[:800],
            "ms": round(ms, 2) if ms is not None else None,
        }
    )
    mark = "PASS" if ok else "FAIL"
    extra = f" ({ms:.0f}ms)" if ms is not None else ""
    print(f"[{mark}] {area} :: {name}{extra} — {detail[:140]}", flush=True)


def client() -> httpx.Client:
    return httpx.Client(base_url=BASE, timeout=180.0)


def wait_alive(timeout_s: float = 120.0) -> bool:
    deadline = time.time() + timeout_s
    while time.time() < deadline:
        try:
            r = httpx.get("http://127.0.0.1:8000/live", timeout=3.0)
            if r.status_code == 200:
                return True
        except Exception:
            pass
        time.sleep(2.0)
    return False


def login(c: httpx.Client, email: str, password: str) -> dict[str, Any] | None:
    t0 = time.perf_counter()
    r = c.post("/auth/login", json={"email": email, "password": password})
    ms = (time.perf_counter() - t0) * 1000
    if r.status_code != 200 or not r.json().get("success"):
        record("auth", f"login {email}", False, r.text[:300], ms)
        return None
    data = r.json()["data"]
    record("auth", f"login {email}", True, f"role={data['user'].get('role')}", ms)
    return data


def _make_text_pdf(text: str) -> bytes:
    """Create a minimal extractable PDF with visible text."""
    # Hand-built PDF with a content stream containing the text.
    # Escape parentheses for PDF string literals.
    safe = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    content = f"BT /F1 12 Tf 50 700 Td ({safe}) Tj ET"
    stream = content.encode("latin-1", errors="replace")
    objects = []
    objects.append(b"1 0 obj<< /Type /Catalog /Pages 2 0 R >>endobj\n")
    objects.append(b"2 0 obj<< /Type /Pages /Kids [3 0 R] /Count 1 >>endobj\n")
    objects.append(
        b"3 0 obj<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>endobj\n"
    )
    objects.append(
        f"4 0 obj<< /Length {len(stream)} >>stream\n".encode("ascii")
        + stream
        + b"\nendstream\nendobj\n"
    )
    objects.append(b"5 0 obj<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>endobj\n")
    out = io.BytesIO()
    out.write(b"%PDF-1.4\n")
    offsets = [0]
    for obj in objects:
        offsets.append(out.tell())
        out.write(obj)
    xref_pos = out.tell()
    out.write(f"xref\n0 {len(offsets)}\n".encode("ascii"))
    out.write(b"0000000000 65535 f \n")
    for off in offsets[1:]:
        out.write(f"{off:010d} 00000 n \n".encode("ascii"))
    out.write(
        f"trailer<< /Size {len(offsets)} /Root 1 0 R >>\nstartxref\n{xref_pos}\n%%EOF\n".encode(
            "ascii"
        )
    )
    return out.getvalue()


def main() -> None:
    metrics: dict[str, Any] = {"started_at": time.time()}
    if not wait_alive():
        print("API not alive")
        return

    with client() as c:
        t0 = time.perf_counter()
        live = httpx.get("http://127.0.0.1:8000/live", timeout=5)
        metrics["live_ms"] = (time.perf_counter() - t0) * 1000
        record("perf", "GET /live", live.status_code == 200, live.text, metrics["live_ms"])

        ready = httpx.get("http://127.0.0.1:8000/ready", timeout=10)
        ready_body = ready.json() if ready.status_code in (200, 503) else {}
        services = (ready_body.get("services") or {})
        # Postgres must be up for this environment; qdrant may be down when using memory fallback
        ready_ok = (
            ready.status_code in (200, 503)
            and services.get("postgres") == "up"
        )
        record(
            "perf",
            "GET /ready",
            ready_ok,
            f"status={ready_body.get('status')} services={services}",
        )
        metrics["ready_services"] = services

        # ---- 1. Auth ----
        admin = login(c, "admin@example.com", "AdminPass123!")
        if not admin:
            print("CRITICAL: admin login failed — aborting")
            dump_report(metrics)
            return

        access = admin["tokens"]["access_token"]
        refresh = admin["tokens"]["refresh_token"]
        headers = {"Authorization": f"Bearer {access}"}

        t0 = time.perf_counter()
        me = c.get("/users/me", headers=headers)
        record(
            "auth",
            "JWT /users/me",
            me.status_code == 200 and me.json().get("success"),
            me.text[:200],
            (time.perf_counter() - t0) * 1000,
        )

        t0 = time.perf_counter()
        ref = c.post("/auth/refresh", json={"refresh_token": refresh})
        ok_ref = ref.status_code == 200 and ref.json().get("success")
        record("auth", "refresh token", ok_ref, ref.text[:200], (time.perf_counter() - t0) * 1000)
        if ok_ref:
            pdata = ref.json()["data"]
            if "access_token" in pdata:
                access = pdata["access_token"]
                refresh = pdata.get("refresh_token") or refresh
            elif "tokens" in pdata:
                access = pdata["tokens"]["access_token"]
                refresh = pdata["tokens"].get("refresh_token") or refresh
            headers = {"Authorization": f"Bearer {access}"}

        t0 = time.perf_counter()
        emb = c.get("/system/embeddings/status", headers=headers)
        record(
            "auth",
            "RBAC admin embeddings status",
            emb.status_code == 200 and emb.json().get("success"),
            json.dumps(emb.json().get("data", {}))[:300],
            (time.perf_counter() - t0) * 1000,
        )
        if emb.status_code == 200:
            metrics["embedding_status"] = emb.json().get("data")

        t0 = time.perf_counter()
        users = c.get("/users", headers=headers, params={"limit": 5})
        record(
            "auth",
            "RBAC admin list users",
            users.status_code == 200 and users.json().get("success"),
            users.text[:200],
            (time.perf_counter() - t0) * 1000,
        )

        uid = uuid.uuid4().hex[:8]
        emp_email = f"e2e_user_{uid}@example.com"
        emp_pass = "UserPass123!"
        t0 = time.perf_counter()
        reg = c.post(
            "/auth/register",
            json={"email": emp_email, "password": emp_pass, "full_name": "E2E User"},
        )
        record(
            "auth",
            "user register",
            reg.status_code in (200, 201) and reg.json().get("success", True),
            reg.text[:250],
            (time.perf_counter() - t0) * 1000,
        )

        user_data = None
        user_headers = None
        if reg.status_code in (200, 201) and reg.json().get("success"):
            if reg.json().get("data", {}).get("tokens"):
                user_data = reg.json()["data"]
                record("auth", f"login {emp_email}", True, "tokens from register", 0)
            else:
                user_data = login(c, emp_email, emp_pass)
        if user_data:
            user_headers = {"Authorization": f"Bearer {user_data['tokens']['access_token']}"}
            t0 = time.perf_counter()
            denied = c.get("/system/embeddings/status", headers=user_headers)
            record(
                "auth",
                "RBAC deny employee admin API",
                denied.status_code in (401, 403) or not denied.json().get("success", True),
                f"status={denied.status_code}",
                (time.perf_counter() - t0) * 1000,
            )

        t0 = time.perf_counter()
        out = c.post("/auth/logout", headers=headers, json={"refresh_token": refresh})
        record(
            "auth",
            "logout",
            out.status_code == 200 and out.json().get("success", True),
            out.text[:200],
            (time.perf_counter() - t0) * 1000,
        )
        admin = login(c, "admin@example.com", "AdminPass123!")
        access = admin["tokens"]["access_token"]
        refresh = admin["tokens"]["refresh_token"]
        headers = {"Authorization": f"Bearer {access}"}

        # ---- 2. Documents ----
        samples: dict[str, tuple[str, str, bytes]] = {
            "txt": (
                "e2e-policy.txt",
                "text/plain",
                b"Remote Work Policy\nEmployees may work remotely up to three days per week.\n"
                b"Managers must approve remote schedules in advance.\n",
            ),
            "csv": ("e2e-data.csv", "text/csv", b"name,role,days\nAlice,Engineer,3\nBob,Manager,2\n"),
            "pdf": (
                "e2e-doc.pdf",
                "application/pdf",
                _make_text_pdf("Deep Learning Basics E2E PDF for retrieval testing"),
            ),
        }
        try:
            from docx import Document as DocxDocument

            dbuf = io.BytesIO()
            doc = DocxDocument()
            doc.add_heading("Vacation Policy", 0)
            doc.add_paragraph("Employees accrue 15 days of vacation per year.")
            doc.save(dbuf)
            samples["docx"] = (
                "e2e-vacation.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                dbuf.getvalue(),
            )
        except Exception as exc:
            record("documents", "create sample DOCX", False, str(exc))

        uploaded_ids: dict[str, str] = {}
        for kind, (fname, mime, payload) in samples.items():
            t0 = time.perf_counter()
            up = c.post(
                "/documents/upload",
                headers=headers,
                files={"file": (fname, payload, mime)},
                data={"visibility": "private", "description": f"e2e {kind}"},
            )
            ms = (time.perf_counter() - t0) * 1000
            ok = up.status_code in (200, 201) and up.json().get("success")
            detail = up.text[:250]
            if ok:
                body = up.json()["data"]
                doc_id = (body.get("document") or body).get("id") or body.get("uuid")
                uploaded_ids[kind] = str(doc_id)
                detail = f"id={doc_id} status={(body.get('document') or body).get('status')}"
            record("documents", f"upload {kind}", ok, detail, ms)

        for kind, doc_id in uploaded_ids.items():
            t0 = time.perf_counter()
            idx = c.post(f"/documents/{doc_id}/index", headers=headers, params={"sync": True})
            ms = (time.perf_counter() - t0) * 1000
            ok = idx.status_code == 200 and idx.json().get("success")
            data = idx.json().get("data") or {}
            chunks = data.get("chunks") or data.get("chunk_count") or 0
            model = data.get("embedding_model") or data.get("model")
            record(
                "documents",
                f"index {kind}",
                ok and data.get("success", True) and int(chunks or 0) >= 0,
                f"chunks={chunks} embed_ms={data.get('embedding_ms')} model={model}",
                ms,
            )
            if data.get("embedding_ms") is not None:
                metrics.setdefault("embed_ms_samples", []).append(data["embedding_ms"])
            # Require at least one chunk for text-bearing docs
            if kind in ("txt", "docx", "csv") and int(chunks or 0) < 1:
                REPORT[-1]["ok"] = False
                REPORT[-1]["detail"] += " | FAIL: expected chunks>=1"

        t0 = time.perf_counter()
        lst = c.get("/documents", headers=headers, params={"limit": 10})
        record(
            "documents",
            "list metadata",
            lst.status_code == 200 and lst.json().get("success"),
            f"total={lst.json().get('data', {}).get('total')}",
            (time.perf_counter() - t0) * 1000,
        )

        # ---- 3. RAG ----
        t0 = time.perf_counter()
        search = c.get("/search", headers=headers, params={"q": "remote work policy", "top_k": 5})
        ms = (time.perf_counter() - t0) * 1000
        hits = (search.json().get("data") or {}).get("hits") or []
        metrics["retrieval_ms"] = ms
        if (search.json().get("data") or {}).get("metrics"):
            metrics["retrieval_engine_metrics"] = search.json()["data"]["metrics"]
        record(
            "rag",
            "semantic search",
            search.status_code == 200 and len(hits) > 0,
            f"hits={len(hits)} top={hits[0].get('filename') if hits else None} score={hits[0].get('confidence') if hits else None}",
            ms,
        )

        t0 = time.perf_counter()
        chat1 = c.post("/chat", headers=headers, json={"message": "What is the remote work policy?"})
        ms = (time.perf_counter() - t0) * 1000
        metrics["llm_response_ms"] = ms
        ok_chat = chat1.status_code == 200 and chat1.json().get("success")
        cdata = chat1.json().get("data") or {}
        cites = cdata.get("citations") or []
        chat_id = cdata.get("chat_id")
        if cdata.get("metrics"):
            metrics["chat_metrics"] = cdata["metrics"]
        record(
            "rag",
            "chat + citations",
            ok_chat and len(cites) > 0,
            f"cites={len(cites)} chat_id={chat_id}",
            ms,
        )

        t0 = time.perf_counter()
        chat2 = c.post(
            "/chat",
            headers=headers,
            json={"message": "How many remote days are allowed?", "chat_id": chat_id},
        )
        record(
            "rag",
            "follow-up question",
            chat2.status_code == 200 and chat2.json().get("success"),
            ((chat2.json().get("data") or {}).get("assistant_message") or {}).get("content", "")[:160],
            (time.perf_counter() - t0) * 1000,
        )

        t0 = time.perf_counter()
        hist = c.get("/chat/history", headers=headers, params={"limit": 10})
        record(
            "rag",
            "conversation history",
            hist.status_code == 200 and hist.json().get("success"),
            f"sessions={hist.json().get('data', {}).get('total')}",
            (time.perf_counter() - t0) * 1000,
        )

        t0 = time.perf_counter()
        try:
            with c.stream(
                "POST",
                "/chat/stream",
                headers=headers,
                json={"message": "Summarize the vacation policy briefly.", "chat_id": chat_id},
            ) as stream:
                chunks = []
                first_ms = None
                for line in stream.iter_lines():
                    if line:
                        if first_ms is None:
                            first_ms = (time.perf_counter() - t0) * 1000
                        chunks.append(line)
                        if len(chunks) > 30:
                            break
                ms = (time.perf_counter() - t0) * 1000
                metrics["streaming_first_token_ms"] = first_ms
                metrics["streaming_batch_ms"] = ms
                record(
                    "rag",
                    "streaming chat",
                    stream.status_code == 200 and len(chunks) > 0,
                    f"lines={len(chunks)} ttfb={first_ms} sample={chunks[0][:80] if chunks else ''}",
                    ms,
                )
        except Exception as exc:
            record("rag", "streaming chat", False, str(exc))

        # ---- 4. OCR ----
        ocr_id = None
        try:
            from PIL import Image, ImageDraw, ImageFont

            img = Image.new("RGB", (400, 120), color=(255, 255, 255))
            ImageDraw.Draw(img).text((10, 40), "INVOICE #E2E-99 Amount Due 42.00", fill=(0, 0, 0))
            ibuf = io.BytesIO()
            img.save(ibuf, format="PNG")
            t0 = time.perf_counter()
            ocr = c.post(
                "/ocr/upload",
                headers=headers,
                files={"file": ("e2e-scan.png", ibuf.getvalue(), "image/png")},
            )
            ok = ocr.status_code in (200, 201) and ocr.json().get("success", True)
            odata = ocr.json().get("data") or {}
            ocr_id = odata.get("id")
            record(
                "ocr",
                "OCR upload/extract",
                ok and bool(odata.get("raw_text") or odata.get("text") or odata.get("id")),
                ocr.text[:300],
                (time.perf_counter() - t0) * 1000,
            )
            # Search OCR store
            t0 = time.perf_counter()
            osearch = c.get("/ocr/search", headers=headers, params={"q": "INVOICE", "limit": 5})
            record(
                "ocr",
                "OCR searchable",
                osearch.status_code == 200 and osearch.json().get("success", True),
                osearch.text[:200],
                (time.perf_counter() - t0) * 1000,
            )
            t0 = time.perf_counter()
            ochat = c.post(
                "/chat",
                headers=headers,
                json={"message": "What invoice number appears in the OCR invoice documents?"},
            )
            ocd = ochat.json().get("data") or {}
            ocites = ocd.get("citations") or []
            ocr_cite = any(
                "ocr" in str(x.get("filename") or "").lower()
                or "invoice" in str(x.get("snippet") or "").lower()
                for x in ocites
            )
            record(
                "ocr",
                "OCR searchable via chat",
                ochat.status_code == 200 and ochat.json().get("success") and (len(ocites) > 0) and ocr_cite,
                f"cites={len(ocites)} ocr_hit={ocr_cite}",
                (time.perf_counter() - t0) * 1000,
            )
        except Exception as exc:
            record("ocr", "OCR upload/extract", False, str(exc))

        # ---- 5. Vision ----
        try:
            from PIL import Image

            vbuf = io.BytesIO()
            Image.new("RGB", (64, 64), color=(30, 144, 255)).save(vbuf, format="PNG")
            t0 = time.perf_counter()
            vis = c.post(
                "/vision/analyze",
                headers=headers,
                files={"file": ("e2e-img.png", vbuf.getvalue(), "image/png")},
            )
            vdata = vis.json().get("data") or {}
            record(
                "vision",
                "image analyze",
                vis.status_code in (200, 201) and vis.json().get("success", True),
                f"caption={str(vdata.get('caption') or '')[:80]} objects={len(vdata.get('objects') or [])}",
                (time.perf_counter() - t0) * 1000,
            )
            t0 = time.perf_counter()
            vh = c.get("/vision/history", headers=headers, params={"limit": 5})
            record(
                "vision",
                "vision stored/history",
                vh.status_code == 200 and vh.json().get("success", True),
                vh.text[:200],
                (time.perf_counter() - t0) * 1000,
            )
        except Exception as exc:
            record("vision", "image analyze", False, str(exc))

        # ---- 6. Meetings ----
        meeting_id = None
        wav = _silent_wav()
        t0 = time.perf_counter()
        meet = c.post(
            "/meetings/upload",
            headers=headers,
            files={"file": ("e2e-meet.wav", wav, "audio/wav")},
            data={"title": "E2E Meeting", "auto_process": "true"},
        )
        mok = meet.status_code in (200, 201) and meet.json().get("success", True)
        mdata = meet.json().get("data") or {}
        meeting_id = mdata.get("id")
        record(
            "meetings",
            "upload+process audio",
            mok,
            f"id={meeting_id} status={mdata.get('status')} speakers={mdata.get('speaker_count')} {meet.text[:180]}",
            (time.perf_counter() - t0) * 1000,
        )
        if meeting_id:
            t0 = time.perf_counter()
            tr = c.get(f"/meetings/{meeting_id}/transcript", headers=headers)
            record(
                "meetings",
                "transcript stored",
                tr.status_code == 200 and tr.json().get("success", True),
                tr.text[:250],
                (time.perf_counter() - t0) * 1000,
            )
            t0 = time.perf_counter()
            mc = c.post(
                f"/meetings/{meeting_id}/chat",
                headers=headers,
                json={"message": "What was discussed in this meeting?"},
            )
            record(
                "meetings",
                "meeting chat",
                mc.status_code == 200 and mc.json().get("success", True),
                mc.text[:250],
                (time.perf_counter() - t0) * 1000,
            )

        # ---- 7. Agents ----
        t0 = time.perf_counter()
        agent = c.post(
            "/agent/run",
            headers=headers,
            json={"goal": "Search company documents for remote work policy", "confirm": True},
        )
        aok = agent.status_code in (200, 201) and agent.json().get("success", True)
        adata = agent.json().get("data") or {}
        tools = adata.get("tool_executions") or []
        record(
            "agents",
            "agent run + tools",
            aok and (len(tools) > 0 or bool(adata.get("answer"))),
            f"tools={len(tools)} answer={str(adata.get('answer') or '')[:100]}",
            (time.perf_counter() - t0) * 1000,
        )
        session_id = adata.get("session_id")
        if session_id:
            t0 = time.perf_counter()
            achat = c.post(
                "/agent/chat",
                headers=headers,
                json={
                    "message": "Remember that remote work allows three days. Confirm from docs.",
                    "session_id": session_id,
                    "confirm": True,
                },
            )
            record(
                "agents",
                "agent memory/follow-up",
                achat.status_code == 200 and achat.json().get("success", True),
                achat.text[:250],
                (time.perf_counter() - t0) * 1000,
            )

        # ---- 8. Analytics ----
        for path, name in (
            ("/analytics/overview", "overview"),
            ("/analytics/rag", "rag metrics"),
            ("/analytics/system", "system metrics"),
            ("/analytics/llm", "llm usage"),
            ("/analytics/agents", "agent metrics"),
            ("/analytics/users", "user usage"),
        ):
            t0 = time.perf_counter()
            a = c.get(path, headers=headers, params={"range": "7d"})
            record(
                "analytics",
                name,
                a.status_code == 200 and a.json().get("success", True),
                f"status={a.status_code} {a.text[:120]}",
                (time.perf_counter() - t0) * 1000,
            )

        # ---- 9. Tenancy ----
        if user_headers and uploaded_ids.get("txt"):
            t0 = time.perf_counter()
            other = c.get(f"/documents/{uploaded_ids['txt']}", headers=user_headers)
            blocked = other.status_code in (401, 403, 404) or not other.json().get("success", True)
            record(
                "tenancy",
                "private doc isolation",
                blocked,
                f"status={other.status_code} body={other.text[:160]}",
                (time.perf_counter() - t0) * 1000,
            )

        t0 = time.perf_counter()
        orgs = c.get("/admin/organizations", headers=headers)
        record(
            "tenancy",
            "admin list organizations",
            orgs.status_code == 200 and orgs.json().get("success"),
            orgs.text[:160],
            (time.perf_counter() - t0) * 1000,
        )

        # ---- 10. Performance ----
        # Process + host metrics (psutil) and analytics snapshot (app)
        try:
            import psutil

            candidates = []
            for p in psutil.process_iter(["pid", "name", "cmdline"]):
                cmd = " ".join(p.info.get("cmdline") or [])
                if "uvicorn" in cmd and "app.main:app" in cmd:
                    candidates.append(psutil.Process(p.info["pid"]))
            if not candidates:
                raise RuntimeError("uvicorn process not found")
            # Prefer the worker with the largest RSS (parent wrapper is tiny)
            proc = max(candidates, key=lambda x: x.memory_info().rss)
            metrics["backend_rss_mb"] = round(proc.memory_info().rss / (1024 * 1024), 1)
            metrics["backend_cpu_percent"] = proc.cpu_percent(interval=0.5)
            metrics["host_ram_percent"] = psutil.virtual_memory().percent
            metrics["host_cpu_percent"] = psutil.cpu_percent(interval=0.2)
            record(
                "perf",
                "backend process resources",
                metrics["backend_rss_mb"] > 0,
                f"rss_mb={metrics['backend_rss_mb']} cpu%={metrics['backend_cpu_percent']} host_ram%={metrics['host_ram_percent']}",
            )
        except Exception as exc:
            record("perf", "backend process resources", False, str(exc))

        t0 = time.perf_counter()
        sysm = c.get("/analytics/system", headers=headers, params={"range": "7d"})
        snap = ((sysm.json().get("data") or {}).get("snapshot") or {})
        metrics["analytics_snapshot"] = snap
        record(
            "perf",
            "analytics vector/system snapshot",
            sysm.status_code == 200
            and sysm.json().get("success", True)
            and int(snap.get("qdrant_points") or 0) > 0,
            f"qdrant_points={snap.get('qdrant_points')} cpu={snap.get('cpu_percent')} ram={snap.get('ram_percent')}",
            (time.perf_counter() - t0) * 1000,
        )

        if metrics.get("embed_ms_samples"):
            samples_ms = metrics["embed_ms_samples"]
            metrics["avg_embedding_ms"] = sum(samples_ms) / len(samples_ms)

    dump_report(metrics)


def _silent_wav(seconds: float = 0.5, rate: int = 16000) -> bytes:
    import wave

    buf = io.BytesIO()
    n = int(seconds * rate)
    with wave.open(buf, "wb") as w:
        w.setnchannels(1)
        w.setsampwidth(2)
        w.setframerate(rate)
        w.writeframes(b"\x00\x00" * n)
    return buf.getvalue()


def dump_report(metrics: dict[str, Any]) -> None:
    passed = [r for r in REPORT if r["ok"]]
    failed = [r for r in REPORT if not r["ok"]]
    out = {
        "summary": {"passed": len(passed), "failed": len(failed), "total": len(REPORT)},
        "metrics": metrics,
        "results": REPORT,
    }
    path = Path(r"C:\scripts\_e2e_report.json")
    path.write_text(json.dumps(out, indent=2), encoding="utf-8")
    # also project copy
    proj = Path(
        r"c:\Users\Whoopit\Desktop\syeda atiqa kanwal\Enterprise AI Knowledge Assistant-20260715T133014Z-1-001\Enterprise AI Knowledge Assistant\scripts\_e2e_report.json"
    )
    proj.parent.mkdir(parents=True, exist_ok=True)
    proj.write_text(json.dumps(out, indent=2), encoding="utf-8")
    print("\n=== SUMMARY ===", flush=True)
    print(f"Passed: {len(passed)}  Failed: {len(failed)}  Total: {len(REPORT)}", flush=True)
    print(f"Report: {path}", flush=True)
    if failed:
        print("\nFAILED:", flush=True)
        for r in failed:
            print(f"  - [{r['area']}] {r['name']}: {r['detail'][:200]}", flush=True)


if __name__ == "__main__":
    main()